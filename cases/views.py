from datetime import date
import tempfile
from docx.opc.exceptions import PackageNotFoundError
from docxtpl import DocxTemplate

from django.conf import settings
from django.contrib import messages
from django.contrib.contenttypes.models import ContentType
from django.db import transaction
from django.db.models import Q, Count, Exists, OuterRef, Sum
from django.db.utils import IntegrityError
from django.http import HttpResponse, HttpResponseRedirect
from django.shortcuts import get_object_or_404
from django.urls import reverse
from django.views.generic import FormView, CreateView, TemplateView
from django.views.generic.detail import DetailView
from django.views.generic.list import ListView
from django.views.generic.edit import UpdateView

from django_import_data.models import FileImporter, ModelImportAttempt
from django_filters.views import FilterView
from django_tables2.export.views import ExportMixin
from django_tables2.views import SingleTableMixin, MultiTableMixin
from watson import search as watson

from audits.filters import FileImporterFilter, ModelImportAttemptFilter
from audits.tables import FileImporterSummaryTable, ModelImportAttemptFailureTable
from utils.coord_utils import coords_to_string
from utils.merge_people import find_similar_people, merge_people
from .forms import LetterTemplateForm, DuplicateCaseForm, CaseForm, PersonForm
from .models import (
    Attachment,
    Case,
    CaseGroup,
    Facility,
    Person,
    PreliminaryCase,
    PreliminaryFacility,
    Structure,
)
from .filters import (
    AttachmentFilter,
    CaseFilter,
    CaseGroupFilter,
    FacilityFilter,
    PersonFilter,
    PreliminaryCaseFilter,
    PreliminaryFacilityFilter,
    StructureFilter,
)
from .tables import (
    AttachmentTable,
    AttachmentDashboardTable,
    CaseExportTable,
    CaseGroupTable,
    CaseTable,
    FacilityExportTable,
    FacilityTable,
    LetterFacilityTable,
    PersonTable,
    PreliminaryCaseExportTable,
    PreliminaryCaseTable,
    PreliminaryFacilityExportTable,
    PreliminaryFacilityTable,
    SearchEntryTable,
    StructureTable,
)
from .kml import (
    facility_as_kml,
    facilities_as_kml,
    case_as_kml,
    cases_as_kml,
    kml_to_string,
)


def get_fields_missing_from_info_tables(context, all_model_fields):
    known_fields = [
        value
        for info_key, info_list in context.items()
        if isinstance(info_list, list)
        for value in info_list
        if info_key.endswith("info")
    ]
    return sorted(f[0] for f in all_model_fields if f[0] not in known_fields)


class FilterTableView(ExportMixin, SingleTableMixin, FilterView):
    table_class = None
    filterset_class = None
    object_list = None
    export_table_class = None
    export_requested = None

    def get(self, request, *args, **kwargs):
        if "show-all" in request.GET:
            self.table_pagination = False

        if "clear" in request.GET:
            return HttpResponseRedirect(
                reverse(f"{self.table_class.Meta.model.__name__.lower()}_index")
            )

        if "csv" in request.GET.get("_export", ""):
            self.export_requested = True
            # Change the value to csv so that django-tables2 understands the
            # export request
            request.GET = request.GET.copy()
            request.GET["_export"] = "csv"
        else:
            self.export_requested = False

        return super().get(request, *args, **kwargs)

    def get_queryset(self):
        return self.table_class.Meta.model.objects.all()

    def get_context_data(self, **kwargs):
        # If there are no query params, then no results are
        # displayed -- but that's not what we want!
        if not self.request.GET:
            # Need this here to avoid a blank table appearing on first load
            self.object_list = self.get_queryset()

        return super().get_context_data(**kwargs)

    def get_table_class(self, **kwargs):
        # If an export has been requested, AND we have a specific table
        # defined for export use, then return that
        if self.export_requested and self.export_table_class:
            return self.export_table_class
        # Otherwise, just act as normal
        else:
            return super().get_table_class(**kwargs)

    def render_to_response(self, context, **response_kwargs):
        if "mass-edit" in self.request.GET and self.object_list:
            verbose_name_plural = (
                self.filterset_class.Meta.model._meta.verbose_name_plural
            )
            messages.info(
                self.request,
                f"Any changes made here will affect {verbose_name_plural}: "
                f"{self.object_list.all()}",
            )
            num_items = self.object_list.count()
            if num_items > 100:
                messages.warning(
                    self.request,
                    f"Are you sure you want to mass edit {num_items} {verbose_name_plural} (it's a lot...)?",
                )

            app = self.filterset_class.Meta.model._meta.app_label
            model = self.filterset_class.Meta.model.__name__.lower()
            instances = ",".join(
                str(item) for item in self.object_list.values_list("id", flat=True)
            )
            return HttpResponseRedirect(
                reverse("massadmin_change_view", args=[app, model, instances])
            )

        return super().render_to_response(context, **response_kwargs)


class CaseGroupDetailView(MultiTableMixin, DetailView):
    model = CaseGroup
    tables = [CaseTable, PreliminaryCaseTable]
    table_pagination = {"per_page": 10}

    def get_tables_data(self):
        case_filter_qs = CaseFilter(
            self.request.GET,
            queryset=self.object.cases.all(),
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        pcase_filter_qs = PreliminaryCaseFilter(
            self.request.GET,
            queryset=self.object.pcases.all(),
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        return [case_filter_qs, pcase_filter_qs]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        context["case_num_ranges"] = self.object.get_case_nums_as_ranges()
        context["pcase_num_ranges"] = self.object.get_pcase_nums_as_ranges()
        return context


class CaseGroupListView(FilterTableView):
    table_class = CaseGroupTable
    filterset_class = CaseGroupFilter
    template_name = "cases/case_group_list.html"

    def get_queryset(self):
        queryset = super().get_queryset()
        # All Cases for "current" CaseGroup that are not completed
        uncompleted = Case.objects.filter(case_groups=OuterRef("id"), completed=False)
        queryset = queryset.annotate(
            num_cases=Count("cases", distinct=True),
            num_pcases=Count("pcases", distinct=True),
            # The inverse of uncompleted will be the completed Cases
            completed=~Exists(uncompleted),
        )
        return queryset


class PreliminaryCaseListView(FilterTableView):
    table_class = PreliminaryCaseTable
    filterset_class = PreliminaryCaseFilter
    export_table_class = PreliminaryCaseExportTable
    template_name = "cases/prelim_case_list.html"


class CaseListView(FilterTableView):
    table_class = CaseTable
    export_table_class = CaseExportTable
    filterset_class = CaseFilter
    template_name = "cases/case_list.html"

    def get_queryset(self):
        queryset = super().get_queryset()
        queryset = Case.objects.annotate_stuff(queryset)
        return queryset

    def get(self, request, *args, **kwargs):
        if "jump_to_case_id" in request.GET:
            case_id = request.GET.get("jump_to_case_id")
            return HttpResponseRedirect(reverse("case_detail", args=[case_id]))

        elif "kml" in request.GET:
            # TODO: Must be a cleaner way to do this
            qs = self.get_filterset(self.filterset_class).qs
            response = HttpResponse(
                kml_to_string(
                    cases_as_kml(qs.filter(facilities__location__isnull=False).all())
                ),
                content_type="application/vnd.google-earth.kml+xml.",
            )
            response["Content-Disposition"] = 'application; filename="nrqz_apps.kml"'
            return response
        else:
            return super(CaseListView, self).get(request, *args, **kwargs)


class PreliminaryFacilityListView(FilterTableView):
    table_class = PreliminaryFacilityTable
    filterset_class = PreliminaryFacilityFilter
    export_table_class = PreliminaryFacilityExportTable
    template_name = "cases/prelim_facility_list.html"

    def get_queryset(self):
        return (
            PreliminaryFacility.objects.annotate_in_nrqz()
            .annotate_azimuth_to_gbt()
            .annotate_distance_to_gbt()
        )


class FacilityListView(FilterTableView):
    table_class = FacilityTable
    filterset_class = FacilityFilter
    export_table_class = FacilityExportTable
    template_name = "cases/facility_list.html"

    def get_queryset(self):
        return (
            Facility.objects.annotate_in_nrqz()
            .annotate_azimuth_to_gbt()
            .annotate_distance_to_gbt()
        )

    def get(self, request, *args, **kwargs):
        if "kml" in request.GET:
            # TODO: Must be a cleaner way to do this
            qs = self.get_filterset(self.filterset_class).qs
            response = HttpResponse(
                kml_to_string(
                    facilities_as_kml(qs.filter(location__isnull=False).all())
                ),
                content_type="application/vnd.google-earth.kml+xml.",
            )
            response[
                "Content-Disposition"
            ] = 'application; filename="nrqz_facilities.kml"'
            return response
        else:
            return super(FacilityListView, self).get(request, *args, **kwargs)


class LetterView(FormView):
    form_class = LetterTemplateForm
    template_name = "cases/concurrence_letter.html"

    def get_initial(self):
        initial = super().get_initial()
        if "facilities" in self.request.GET:
            initial.update({"facilities": self.request.GET.getlist("facilities")})

        if "cases" in self.request.GET:
            initial.update({"cases": self.request.GET.getlist("cases")})

        if "template" in self.request.GET:
            initial.update({"template": self.request.GET["template"]})

        return initial

    @staticmethod
    def get_letter_context(post_dict):
        cases = post_dict["cases"]
        facilities = post_dict["facilities"]

        if cases.count() != 1:
            raise ValueError(
                f"There should only be one unique case! Got {cases.count()}!"
            )

        case = cases.first()

        letter_context = {"case": case, "facilities": facilities}

        letter_context["generation_date"] = date.today().strftime("%B %d, %Y")
        letter_context["nrqz_ids"] = ", ".join(
            facilities.filter(nrqz_id__isnull=False).values_list("nrqz_id", flat=True)
        )
        table = LetterFacilityTable(data=facilities)
        letter_context["facilities_table"] = table

        return letter_context

    def form_valid(self, form):
        # We make a temp file...
        letter_context = self.get_letter_context(form.cleaned_data)
        letter_template = form.cleaned_data["template"]
        try:
            dt = DocxTemplate(letter_template.path)
        except PackageNotFoundError as error:
            messages.error(
                self.request,
                f"Letter template at {letter_template.path} does not exist! "
                "You will need to resolve this manually.",
            )
            return HttpResponseRedirect(reverse("letters"))

        with tempfile.NamedTemporaryFile() as fp:
            dt.render(letter_context)
            # ...write the converted document to it...
            dt.save(fp.name)
            # ...and then read it into memory
            docx = fp.read()

        # Generate the filename based on the case number(s)
        filename = f"{letter_context['case'].case_num}_letter.docx"

        # And serve the document
        # TODO: Large files will probably cause issues here... will need to set up streaming if
        # this happens
        response = HttpResponse(
            docx,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'application; filename="{filename}"'
        return response


class PreliminaryCaseDetailView(MultiTableMixin, DetailView):
    model = PreliminaryCase
    tables = [
        PreliminaryFacilityTable,
        AttachmentTable,
        PreliminaryCaseTable,
        CaseTable,
    ]
    table_pagination = {"per_page": 10}

    def get_tables_data(self):
        pfacility_filter_qs = PreliminaryFacilityFilter(
            self.request.GET,
            queryset=self.object.pfacilities.all(),
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        attachment_filter_qs = AttachmentFilter(
            self.request.GET,
            queryset=self.object.attachments.all(),
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        pcase_filter_qs = PreliminaryCaseFilter(
            self.request.GET,
            queryset=self.object.related_prelim_cases,
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        case_filter_qs = CaseFilter(
            self.request.GET,
            queryset=self.object.related_cases,
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        return [
            pfacility_filter_qs,
            attachment_filter_qs,
            pcase_filter_qs,
            case_filter_qs,
        ]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["status_info"] = ["completed_on"]
        context["application_info"] = ["radio_service", "num_freqs", "num_sites"]
        if settings.DEBUG:
            context["unsorted_info"] = get_fields_missing_from_info_tables(
                context, self.object.all_fields()
            )
        return context


class CaseDetailView(MultiTableMixin, DetailView):
    model = Case
    tables = [
        FacilityTable,
        AttachmentTable,
        CaseTable,
        PreliminaryCaseTable,
        FileImporterSummaryTable,
        ModelImportAttemptFailureTable,
    ]
    table_pagination = {"per_page": 10}

    def get_tables_data(self):
        facility_filter_qs = FacilityFilter(
            self.request.GET,
            queryset=self.object.facilities.all().annotate_azimuth_to_gbt(),
            form_helper_kwargs={"form_class": "collapse"},
        ).qs

        attachment_filter_qs = AttachmentFilter(
            self.request.GET,
            queryset=self.object.attachments,
            form_helper_kwargs={"form_class": "collapse"},
        ).qs

        case_filter_qs = CaseFilter(
            self.request.GET,
            queryset=self.object.related_cases,
            form_helper_kwargs={"form_class": "collapse"},
        ).qs

        pcase_filter_qs = PreliminaryCaseFilter(
            self.request.GET,
            queryset=self.object.related_prelim_cases,
            form_helper_kwargs={"form_class": "collapse"},
        ).qs

        if self.object.model_import_attempt:
            mia_failures_filter_qs = ModelImportAttemptFilter(
                self.request.GET,
                queryset=ModelImportAttempt.objects.all()
                # .annotate_is_latest()
                .filter(
                    id__in=(
                        self.object.model_import_attempt.model_importer.row_data.model_importers.values(
                            "model_import_attempts"
                        )
                    ),
                    status=ModelImportAttempt.STATUSES.rejected.db_value,
                    # is_latest=True,
                ),
                form_helper_kwargs={"form_class": "collapse"},
            ).qs
        else:
            mia_failures_filter_qs = ModelImportAttempt.objects.none()

        # Get the File Importers that contain the MIAs used to create this Case's Facilities
        related_file_importers = FileImporter.objects.prefetch_related(
            "file_import_attempts__row_datas__model_importers__model_import_attempts"
        ).filter(
            file_import_attempts__row_datas__model_importers__model_import_attempts__in=(
                facility_filter_qs.values("model_import_attempt")
            )
        )
        fi_filter_qs = FileImporterFilter(
            self.request.GET,
            queryset=related_file_importers.order_by("id")
            .annotate_current_status()
            .distinct(),
        ).qs
        return [
            facility_filter_qs,
            attachment_filter_qs,
            case_filter_qs,
            pcase_filter_qs,
            fi_filter_qs,
            mia_failures_filter_qs,
        ]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        si_done = self.object.get_si_done()
        context["status_info"] = [
            "shutdown",
            (
                "Completed On",
                self.object.completed_on if self.object.completed_on else None,
                "",
            ),
            "si_waived",
            "si",
            (
                "SI Done",
                si_done if si_done else None,
                "This is the 'SI Done' that is derived from "
                "the per-facility inspection information",
            ),
            (
                "Case-level SI Done",
                self.object.original_si_done,
                "This is the 'SI Done' that is stored per-case, without consideration "
                "of per-facility inspection information",
            ),
            "is_federal",
        ]
        context["application_info"] = [
            "radio_service",
            "call_sign",
            "freq_coord",
            "fcc_file_num",
            ("# Facilities Evaluated", self.object.num_facilities_evaluated, ""),
            "num_sites",
            "num_freqs",
            "num_outside",
            ("Meets ERPd Limit", self.object.get_meets_erpd_limit, ""),
            "agency_num",
        ]
        context["sgrs_info"] = [
            "sgrs_notify",
            (
                "SGRS Responded On",
                self.object.sgrs_responded_on
                if self.object.sgrs_responded_on
                else None,
                "",
            ),
            "sgrs_service_num",
            ("SGRS Approval", self.object.get_sgrs_approval, ""),
        ]
        if settings.DEBUG:
            context["unsorted_info"] = get_fields_missing_from_info_tables(
                context, self.object.all_fields()
            )
        context["duplicate_case_form"] = DuplicateCaseForm()
        return context

    def as_kml(self):
        return case_as_kml(self.object)


class BaseFacilityDetailView(DetailView):
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["topography_info"] = [
            "site_name",
            "call_sign",
            "location_description",
            "location",
            "amsl",
            "survey_1a",
            "survey_2c",
            (
                "Inside NRQZ?",
                self.object.get_in_nrqz(),
                "Indicates whether the facility is inside the boundaries of the NRQZ",
            ),
        ]

        context["path_data_interpolation_info"] = [
            "topo_4_point",
            "topo_12_point",
            "propagation_model",
        ]
        # TODO:
        context["usgs_dataset_info"] = ["usgs_dataset"]

        context["emissions_info"] = ["bandwidth", "emissions"]

        context["path_attenuation_info"] = [
            "nrao_diff",
            "nrao_space",
            "nrao_tropo",
            "tpa",
            "distance_to_first_obstacle",
            "height_of_first_obstacle",
            (
                "Azimuth Bearing (derived)",
                f"{self.object.get_azimuth_to_gbt():.3f}°"
                if self.object.location
                else None,
                "Azimuth bearing to GBT in degrees",
            ),
            "dominant_path",
            ("Propagation Study", self.object.get_prop_study_as_link(), ""),
        ]

        context["sgrs_info"] = [
            "sgrs_approval",
            "sgrs_responded_on",
            "sgrs_work_order_num",
        ]

        return context


class PreliminaryFacilityDetailView(BaseFacilityDetailView):
    model = PreliminaryFacility

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["location_info"] = ["location", "amsl", "agl"]
        context["topography_info"] = [
            item for item in context["topography_info"] if item != "call_sign"
        ]
        context["antenna_info"] = [
            "agl",
            "antenna_model_number",
            (
                "Azimuth Bearing (derived)",
                f"{self.object.get_azimuth_to_gbt():.3f}°"
                if self.object.location
                else None,
                "Azimuth bearing to GBT in degrees",
            ),
        ]
        context["transmitter_info"] = [
            "freq_low",
            "freq_high",
            "power_density_limit",
            # "tx_per_sector",
            # "max_tx_power",
            # "num_tx_per_facility",
            # "max_erp_per_tx",
        ]
        if settings.DEBUG:
            context["unsorted_info"] = get_fields_missing_from_info_tables(
                context, self.object.all_fields()
            )
        context["federal_info"] = [
            ("Is Federal?", self.object.pcase.is_federal, ""),
            "agency_num",
            "s367",
        ]
        return context


class FacilityDetailView(MultiTableMixin, BaseFacilityDetailView):
    model = Facility
    tables = [AttachmentTable]
    table_pagination = {"per_page": 10}

    def get_tables_data(self):
        attachment_filter_qs = AttachmentFilter(
            self.request.GET,
            queryset=self.object.attachments.all(),
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        return [attachment_filter_qs]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["antenna_info"] = [
            "agl",
            "antenna_model_number",
            "antenna_gain",
            # az deg true
            "main_beam_orientation",
            "mechanical_downtilt",
            "electrical_downtilt",
        ]
        context["transmitter_info"] = [
            "freq_low",
            "freq_high",
            "power_density_limit",
            "tx_per_sector",
            "max_tx_power",
            "num_tx_per_facility",
            "requested_max_erp_per_tx",
        ]
        if settings.DEBUG:
            context["unsorted_info"] = get_fields_missing_from_info_tables(
                context, self.object.all_fields()
            )

        context["analysis_results_info"] = [
            "meets_erpd_limit",
            "nrao_aerpd",
            "nrao_aerpd_analog",
            "nrao_aerpd_gsm",
            "nrao_aerpd_cdma",
            "nrao_aerpd_cdma2000",
        ]
        context["federal_info"] = [
            ("Is Federal?", self.object.case.is_federal, ""),
            "agency_num",
            "s367",
        ]
        context["fcc_info"] = ["fcc_file_number"]
        return context

    def as_kml(self):
        facility_as_kml(self.object)


class AttachmentListView(FilterTableView):
    table_class = AttachmentTable
    filterset_class = AttachmentFilter
    template_name = "cases/attachment_list.html"

    def post(self, request, *args, **kwargs):
        if request.POST.get("refresh_from_filesystem", None):
            report = self.table_class.Meta.model.objects.all().refresh_from_filesystem()
            report_counts = {key: len(values) for key, values in report.items()}
            messages.success(
                request,
                f"Successfully refreshed all attachments from filesystem: {report_counts}",
            )
            return HttpResponseRedirect(reverse("attachment_index"))

        return super().post(request, *args, **kwargs)


class AttachmentDetailView(MultiTableMixin, DetailView):
    model = Attachment
    tables = [CaseTable, PreliminaryCaseTable, FacilityTable, PreliminaryFacilityTable]
    table_pagination = {"per_page": 10}

    def get_tables_data(self):
        case_filter_qs = CaseFilter(
            self.request.GET,
            queryset=self.object.cases.all(),
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        pcase_filter_qs = PreliminaryCaseFilter(
            self.request.GET,
            queryset=self.object.prelim_cases.all(),
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        facility_filter_qs = FacilityFilter(
            self.request.GET,
            queryset=self.object.facilities.all(),
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        pfacility_filter_qs = PreliminaryFacilityFilter(
            self.request.GET,
            queryset=self.object.prelim_facilities.all(),
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        return [
            case_filter_qs,
            pcase_filter_qs,
            facility_filter_qs,
            pfacility_filter_qs,
        ]

    def get_object(self, *args, **kwargs):
        instance = super().get_object(*args, **kwargs)
        # Refresh the file info from disk
        status = instance.refresh_from_filesystem()
        if status == "changed":
            messages.warning(
                self.request,
                "Attachment contents have changed since last checked! Stored "
                "hash has been updated",
            )
        return instance


class PersonListView(FilterTableView):
    table_class = PersonTable
    filterset_class = PersonFilter
    template_name = "cases/person_list.html"


class PersonDetailView(MultiTableMixin, DetailView):
    tables = [CaseTable, PreliminaryCaseTable, PersonTable]
    model = Person
    table_pagination = {"per_page": 10}

    def get_tables_data(self):
        case_filter_qs = CaseFilter(
            self.request.GET,
            queryset=Case.objects.filter(
                Q(applicant=self.object) | Q(contact=self.object)
            ),
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        pcase_filter_qs = PreliminaryCaseFilter(
            self.request.GET,
            queryset=PreliminaryCase.objects.filter(
                Q(applicant=self.object) | Q(contact=self.object)
            ),
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        person_filter_qs = PersonFilter(
            self.request.GET,
            queryset=find_similar_people(self.object),
            form_helper_kwargs={"form_class": "collapse"},
        ).qs
        return [case_filter_qs, pcase_filter_qs, person_filter_qs]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        context["applicant_info"] = [
            # "cases",
            "name",
            "phone",
            "fax",
            (
                "email",
                f'<a href="mailto:{self.object.email}">{self.object.email}</a>',
                "",
            ),
            "street",
            "city",
            "county",
            "state",
            "zipcode",
        ]

        return context


def merge_similar_people(request, pk):
    """Given a Person ID, find people similar to them and merge them all together

    Person with given ID will be kept; others will be merged
    """

    person = get_object_or_404(Person, id=pk)
    similar_people = find_similar_people(person).exclude(id=person.id)
    if similar_people:
        with transaction.atomic():
            merge_people(person_to_keep=person, people_to_merge=similar_people)
        messages.success(
            request,
            f"Successfully merged {person.merge_info.num_instances_merged - 1} "
            f"people into {person.name}",
        )
    else:
        messages.warning(request, f"There are no similar people to merge!")
    return HttpResponseRedirect(person.get_absolute_url())


class StructureListView(FilterTableView):
    table_class = StructureTable
    filterset_class = StructureFilter
    template_name = "cases/structure_list.html"

    def get_queryset(self):
        queryset = super().get_queryset()
        queryset = queryset.annotate(
            num_facilities=Count("facilities"),
            num_cases=Count("facilities__case", distinct=True),
        )
        return queryset


class StructureDetailView(DetailView):
    model = Structure

    def __init__(self, *args, **kwargs):
        super(StructureDetailView, self).__init__(*args, **kwargs)
        self.facility_filter = None

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        context["info"] = [
            "asr",
            "file_num",
            (
                "Location",
                coords_to_string(
                    latitude=self.object.location.y,
                    longitude=self.object.location.x,
                    concise=True,
                ),
                "",
            ),
            "faa_circ_num",
            "faa_study_num",
            "issue_date",
            "height",
        ]

        if not self.facility_filter:
            self.facility_filter = FacilityFilter(
                self.request.GET,
                queryset=self.object.facilities.all(),
                form_helper_kwargs={"form_class": "collapse"},
            )
            context["facility_filter"] = self.facility_filter

        if "facility_table" not in context:
            table = FacilityTable(data=self.facility_filter.qs)
            table.paginate(page=self.request.GET.get("page", 1), per_page=10)
            context["facility_table"] = table

        return context


class SearchView(MultiTableMixin, ListView):
    tables = [SearchEntryTable, SearchEntryTable, SearchEntryTable]
    table_pagination = {"per_page": 10}
    template_name = "cases/search_results.html"

    def get_tables_data(self):
        for table in self.tables:
            table.query = self.query

        data = [
            self.object_list.filter(
                content_type=ContentType.objects.get_for_model(model)
            )
            for model in [Case, Facility, Person]
        ]
        return data

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["query"] = self.query
        context["tables_with_model_names"] = zip(
            ["Case", "Facility", "Person"], context["tables"]
        )
        return context

    def get(self, request, *args, **kwargs):
        self.query = request.GET.get("q", None)
        self.object_list = watson.search(self.query)
        context = self.get_context_data()
        return self.render_to_response(context)


def facility_as_kml_view(request, pk):
    facility = Facility.objects.get(id=pk)
    response = HttpResponse(
        facility.as_kml(), content_type="application/vnd.google-earth.kml+xml."
    )
    response["Content-Disposition"] = f'application; filename="{facility.nrqz_id}.kml"'
    return response


def case_as_kml_view(request, pk):
    case = Case.objects.get(id=pk)
    response = HttpResponse(
        case.as_kml(), content_type="application/vnd.google-earth.kml+xml."
    )
    response["Content-Disposition"] = f'application; filename="{case.case_num}.kml"'
    return response


@transaction.atomic
def duplicate_case(request, case_num):
    case = get_object_or_404(Case, case_num=case_num)
    num_duplicates = int(request.POST.get("num_duplicates"))

    successful_duplications = []
    failed_duplications = []
    for __ in range(num_duplicates):
        case.id = None
        case.case_num += 1
        try:
            with transaction.atomic():
                case.save()
        except IntegrityError as error:
            if "duplicate" in error.args[0]:
                failed_duplications.append(case.case_num)
            else:
                raise
        else:
            successful_duplications.append(case.case_num)

    if failed_duplications:
        transaction.set_rollback(True)
        if len(failed_duplications) > 1:
            message_text = (
                f"Cases {failed_duplications} already exist! No new cases created"
            )
        elif len(failed_duplications) == 1:
            message_text = (
                f"Case {failed_duplications[0]} already exists! No new case created"
            )
        else:
            raise ValueError("This shouldn't be possible...")

        messages.error(request, message_text)
        return HttpResponseRedirect(reverse("case_detail", args=[case_num]))

    if len(successful_duplications) > 1:
        message_text = f"Cases {successful_duplications} have been successfully duplicated from case {case_num}!"

    elif len(successful_duplications) == 1:
        message_text = f"Case {successful_duplications[0]} has been successfully duplicated from case {case_num}!"

    else:
        raise ValueError("This shouldn't be possible...")

    messages.success(request, message_text)
    return HttpResponseRedirect(reverse("case_detail", args=[case.case_num]))


class CaseCreateView(CreateView):
    form_class = CaseForm
    template_name = "cases/case_create.html"


class CaseUpdateView(UpdateView):
    model = Case
    form_class = CaseForm
    template_name = "cases/case_update.html"


class PersonCreateView(CreateView):
    form_class = PersonForm
    template_name = "cases/person_create.html"


class PersonUpdateView(UpdateView):
    model = Person
    form_class = PersonForm
    template_name = "cases/person_update.html"


class AttachmentDashboard(SingleTableMixin, TemplateView):
    table_class = AttachmentDashboardTable
    template_name = "audits/attachment_dashboard.html"
    table_pagination = {"per_page": 10}

    def get_queryset(self):
        return Attachment.objects.filter(hash_on_disk__isnull=True, is_active=True)

    def post(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        affect_all = request.POST.get("all", None)
        attachments_to_affect = queryset.filter(is_active=True)
        if not affect_all:
            attachment_ids = request.POST.getlist("check", None)
            attachments_to_affect = attachments_to_affect.filter(id__in=attachment_ids)

        num_attachments_to_affect = attachments_to_affect.count()
        if num_attachments_to_affect:
            if request.POST.get("submit_refresh", None):
                report = attachments_to_affect.refresh_from_filesystem()
                report_counts = {key: len(values) for key, values in report.items()}
                messages.success(
                    request,
                    f"Successfully refreshed {num_attachments_to_affect} attachments from filesystem: {report_counts}",
                )
            elif request.POST.get("submit_deactivate", None):
                # Get count here, since QS will be empty soon
                # Convert to string here, since this QS will be empty soon. We rely
                # on Django to concatenate the values list string to a reasonable length,
                # so we don't have to worry about doing it ourselves
                summary_str = str(attachments_to_affect.values_list("id", flat=True))
                # Perform the actual deactivation
                attachments_to_affect.update(is_active=False)
                if num_attachments_to_affect:
                    messages.success(
                        request,
                        f"Successfully deactivated {num_attachments_to_affect} Attachments: {summary_str}",
                    )
        else:
            messages.warning(request, "No Attachments selected")

        return HttpResponseRedirect(reverse("attachment_dashboard"))
